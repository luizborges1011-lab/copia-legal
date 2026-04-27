#!/usr/bin/env python3
"""
Gerador de Alteração Contratual (LTDA e LTDA Unipessoal)
Gera o instrumento de alteração + contrato social consolidado.
"""

from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from num2words import num2words

# ---------------------------------------------------------------------------
# Importa utilitários compartilhados com gerar_contrato.py
# ---------------------------------------------------------------------------
from gerar_contrato import (
    FONT_NAME, FONT_SIZE, INDENT_NORMAL, INDENT_SOCIO, ESTADOS_NOMES,
    SPC_NORMAL, SPC_CLAUSULA, SPC_CNAE, SPC_CAPITAL,
    SPC_DATA, SPC_ABERTURA, SPC_ASSINATURA, MESES_PT,
    formatar_cpf, formatar_cep, formatar_valor_reais, valor_por_extenso,
    cotas_por_extenso, formatar_data_por_extenso, inferir_genero,
    normalizar_profissao, formatar_endereco, estado_civil_texto,
    formatar_documento, socio_qualificado, title_case,
    set_paragraph_format, add_run, set_table_width, set_col_width,
    format_cell_text, remover_bordas_tabela,
    get_integralizacoes_socio, texto_tipo_integralizacao,
    resolver_forma_integralizacao,
)
from db import get_texto_alteracao


def _t(codigo: str, fallback: str = "") -> str:
    """Lê texto editável do banco para cláusulas de alteração."""
    val = get_texto_alteracao(codigo)
    return val if val else fallback

# Dimensões da tabela de cotas (definidas localmente em gerar_contrato.gerar_contrato)
COL_WIDTHS  = [5118, 1138, 1274, 1844]
TABLE_WIDTH = 9374

# ---------------------------------------------------------------------------
ROMANOS = [
    "I","II","III","IV","V","VI","VII","VIII","IX","X",
    "XI","XII","XIII","XIV","XV","XVI","XVII","XVIII","XIX","XX",
    "XXI","XXII","XXIII","XXIV","XXV"
]

ORDINALS_F = {
    1:"1ª", 2:"2ª", 3:"3ª", 4:"4ª", 5:"5ª",
    6:"6ª", 7:"7ª", 8:"8ª", 9:"9ª", 10:"10ª"
}

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def romano(n: int) -> str:
    return ROMANOS[n - 1] if 1 <= n <= len(ROMANOS) else str(n)


def _texto_integralizacao_partes(cs: dict, valor_sub: float) -> str:
    """Gera texto de integralização a partir de integ_partes (novo) ou fallback integ_tipo (legado)."""
    partes = [p for p in cs.get("integ_partes", []) if float(p.get("valor", 0)) > 0.01]

    if not partes:
        # Fallback para estrutura antiga (integ_tipo + integ_desc)
        tipo = cs.get("integ_tipo", "moeda")
        desc = cs.get("integ_desc", "")
        if tipo == "bens_moveis":
            return f"em bens móveis{f' ({desc})' if desc else ''}"
        if tipo == "bens_imoveis":
            return f"em bem imóvel{f' ({desc})' if desc else ''}"
        return "em moeda corrente nacional"

    def _parte_str(p: dict) -> str:
        v   = formatar_valor_reais(p["valor"])
        ext = valor_por_extenso(p["valor"])
        if p["tipo"] == "bens_moveis":
            desc = f" ({p['descricao']})" if p.get("descricao") else ""
            return f"R$ {v} ({ext}) em bens móveis{desc}"
        if p["tipo"] == "bens_imoveis":
            desc = f" ({p['descricao']})" if p.get("descricao") else ""
            return f"R$ {v} ({ext}) em bem imóvel{desc}"
        return f"R$ {v} ({ext}) em moeda corrente nacional"

    if len(partes) == 1:
        p = partes[0]
        if p["tipo"] == "bens_moveis":
            desc = f" ({p['descricao']})" if p.get("descricao") else ""
            return f"em bens móveis{desc}"
        if p["tipo"] == "bens_imoveis":
            desc = f" ({p['descricao']})" if p.get("descricao") else ""
            return f"em bem imóvel{desc}"
        return "em moeda corrente nacional"

    textos = [_parte_str(p) for p in partes]
    return "sendo " + ", ".join(textos[:-1]) + " e " + textos[-1]


def ordinal(n: int, fem=True) -> str:
    sufixos_f = {1:"primeira",2:"segunda",3:"terceira",4:"quarta",5:"quinta",
                 6:"sexta",7:"sétima",8:"oitava",9:"nona",10:"décima"}
    sufixos_m = {1:"primeiro",2:"segundo",3:"terceiro",4:"quarto",5:"quinto",
                 6:"sexto",7:"sétimo",8:"oitavo",9:"nono",10:"décimo"}
    dic = sufixos_f if fem else sufixos_m
    return dic.get(n, f"{n}ª" if fem else f"{n}º")


def adicionar_cabecalho_alteracao(doc, razao_social: str, num_alt: int,
                                   cnpj: str = "", nire: str = ""):
    section = doc.sections[0]
    header  = section.header
    for p in header.paragraphs:
        p.clear()

    def _hpar(par):
        par.paragraph_format.space_before = 0
        par.paragraph_format.space_after  = 0
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p1 = header.paragraphs[0]
    _hpar(p1)
    add_run(p1, f"{ordinal(num_alt).upper()} ALTERAÇÃO DO CONTRATO SOCIAL", bold=True)

    p2 = header.add_paragraph()
    _hpar(p2)
    add_run(p2, razao_social.upper(), bold=True)

    if cnpj:
        p3 = header.add_paragraph()
        _hpar(p3)
        add_run(p3, f"CNPJ: {cnpj}")

    if nire:
        p4 = header.add_paragraph()
        _hpar(p4)
        add_run(p4, f"NIRE: {nire}")


def adicionar_cabecalho_consolidado(doc, razao_social: str, cnpj: str, nire: str):
    """Cabeçalho da seção consolidada (sem header de página — vai inline)."""
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         left_indent=0, space_after=SPC_CLAUSULA)
    add_run(p, "CONTRATO SOCIAL CONSOLIDADO", bold=True)

    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         left_indent=0, space_after=SPC_CLAUSULA)
    add_run(p, razao_social.upper(), bold=True)

    if cnpj:
        p = doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             left_indent=0, space_after=SPC_CLAUSULA)
        add_run(p, f"CNPJ: {cnpj}")

    if nire:
        p = doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             left_indent=0, space_after=SPC_NORMAL)
        add_run(p, f"NIRE: {nire}")


def bloco_assinatura_cell(cell, nome: str, papel: str):
    from docx.shared import Pt as _Pt
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = 0
    p.paragraph_format.space_after  = 0
    add_run(p, "_____________________________")
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = 0
    p2.paragraph_format.space_after  = 0
    add_run(p2, nome.upper(), bold=True)
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = 0
    p3.paragraph_format.space_after  = 0
    add_run(p3, papel, bold=True)


def papel_socio(s: dict, tipo: str = "normal") -> str:
    """
    tipo: "normal" | "ingressante" | "ex"
    """
    gen = inferir_genero(s)
    adm = s.get("administrador", False)
    if tipo == "ingressante":
        return "Sócia/Ingressante" if gen == "f" else "Sócio/Ingressante"
    if tipo == "ex":
        if gen == "f":
            return "Ex-Sócia/Administradora" if adm else "Ex-Sócia"
        return "Ex-Sócio/Administrador" if adm else "Ex-Sócio"
    # normal
    if gen == "f":
        return "Sócia/Administradora" if adm else "Sócia"
    return "Sócio/Administrador" if adm else "Sócio"


def gerar_tabela_cotas(doc, socios, total_cotas=None):
    """Gera a tabela de distribuição de cotas."""
    if total_cotas is None:
        total_cotas = sum(int(s.get("quantidadeCotas", 0)) for s in socios)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_table_width(table, TABLE_WIDTH)

    headers = ["Sócios", "%", "Quotas", "Valor em R$"]
    for i, (cell, txt) in enumerate(zip(table.rows[0].cells, headers)):
        set_col_width(cell, COL_WIDTHS[i])
        format_cell_text(cell, txt, bold=True)

    total_pct   = 0.0
    total_valor = 0.0

    for socio in socios:
        qtd = int(socio.get("quantidadeCotas", 0))
        val = qtd * float(socio.get("valorUnitarioCota", 1))
        pct = (qtd / total_cotas * 100) if total_cotas else 0
        total_pct   += pct
        total_valor += val

        row = table.add_row().cells
        for i, (cell, txt) in enumerate(zip(row, [
            socio["nome"].upper(),
            f"{pct:.2f}".replace(".", ","),
            f"{qtd:,}".replace(",", "."),
            formatar_valor_reais(val)
        ])):
            set_col_width(cell, COL_WIDTHS[i])
            format_cell_text(cell, txt)

    row = table.add_row().cells
    for i, (cell, txt) in enumerate(zip(row, [
        "TOTAL:",
        f"{total_pct:.2f}".replace(".", ","),
        f"{total_cotas:,}".replace(",", "."),
        formatar_valor_reais(total_valor)
    ])):
        set_col_width(cell, COL_WIDTHS[i])
        format_cell_text(cell, txt, bold=True)


# ---------------------------------------------------------------------------
# Gera o instrumento de alteração
# ---------------------------------------------------------------------------

def gerar_instrumento(doc, dados: dict, empresa_atual: dict, empresa_nova: dict,
                       socios_atuais: list, socios_novos: list, alteracoes: dict,
                       num_alt: int, data_fmt: str, cidade_foro: str, estado_foro: str):

    # --- Abertura ---
    p = doc.add_paragraph()
    set_paragraph_format(p, space_after=SPC_ABERTURA, left_indent=INDENT_NORMAL)
    qtd  = len(socios_atuais)
    gens = [inferir_genero(s) for s in socios_atuais]
    if qtd == 1:
        intro = ("O sócio abaixo identificado e qualificado:"
                 if gens[0] == "m" else "A sócia abaixo identificada e qualificada:")
    elif all(g == "f" for g in gens):
        intro = "As sócias abaixo identificadas e qualificadas:"
    else:
        intro = "Os sócios abaixo identificados e qualificados:"
    add_run(p, intro)

    # Qualificação dos sócios atuais
    for s in socios_atuais:
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_SOCIO, space_after=SPC_NORMAL)
        add_run(p, s["nome"].upper(), bold=True)
        add_run(p, ", " + socio_qualificado(s) + ";")

    # Preâmbulo: "Únicos sócios... registrada na Junta Comercial... resolvem alterar e consolidar..."
    sede_atual   = formatar_endereco(empresa_atual.get("enderecoComercial", {}))
    cnpj         = empresa_atual.get("cnpj", "")
    nire         = empresa_atual.get("nire", "")
    estado_uf    = empresa_atual.get("enderecoComercial", {}).get("estado", "PR").strip().upper()
    estado_junta = ESTADOS_NOMES.get(estado_uf, estado_uf)

    if qtd == 1:
        comp_txt = "Única sócia componente" if gens[0] == "f" else "Único sócio componente"
    elif all(g == "f" for g in gens):
        comp_txt = "Únicas sócias componentes"
    else:
        comp_txt = "Únicos sócios componentes"

    resolve_txt = "resolve" if qtd == 1 else "resolvem"

    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=INDENT_SOCIO, space_after=SPC_NORMAL)
    add_run(p, f"{comp_txt} da sociedade empresária limitada, que gira sob o nome de ")
    add_run(p, empresa_atual.get("razaoSocial", "").upper(), bold=True)
    add_run(p, f", com sede na {sede_atual}")
    if cnpj:
        add_run(p, f", inscrita no CNPJ: {cnpj}")
    if nire:
        add_run(p, f", registrada na Junta Comercial do {estado_junta} sob n.º {nire}")
    resolve_body = _t("alt_preambulo_resolve",
        "resolvem alterar e consolidar o contrato social primitivo e demais alterações, "
        "mediante as condições estabelecidas nas cláusulas seguintes:")
    add_run(p, f", {resolve_txt} {resolve_body}")

    # --- Cláusulas de alteração ---
    num_clausula = 1

    # NOME EMPRESARIAL
    if alteracoes.get("nome_empresarial", {}).get("ativo"):
        novo_nome = alteracoes["nome_empresarial"].get("novo", "").upper()
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CLAUSULA)
        add_run(p, f"CLÁUSULA {romano(num_clausula)} – DO NOME EMPRESARIAL", bold=True)
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)
        add_run(p, "O nome empresarial da sociedade passa a ser: ")
        add_run(p, novo_nome, bold=True)
        add_run(p, ", mantidos os demais dados cadastrais e registros junto às repartições competentes.")
        num_clausula += 1

    # ENDEREÇO
    if alteracoes.get("endereco", {}).get("ativo"):
        novo_end = alteracoes["endereco"]["novo"]
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CLAUSULA)
        add_run(p, f"CLÁUSULA {romano(num_clausula)} - DA SEDE (art. 997, II, CC)", bold=True)
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)
        add_run(p, "O endereço da Sede passa a ser na: ")
        add_run(p, formatar_endereco(novo_end) + ".", bold=True)
        num_clausula += 1

    # Helper para lookup de sócio atual por nome
    def _socio_atual(nome: str):
        for s in socios_atuais:
            if s.get("nome","").upper() == nome.upper():
                return s
        return {}

    # Helper para gerar cláusula de ingresso + sub-rogação
    def _clausula_ingresso(s: dict, gerar_subroga: bool = True):
        nonlocal num_clausula
        gen = inferir_genero(s)
        titulo_ing = "INGRESSO DE SÓCIO" if gen == "m" else "INGRESSO DE SÓCIA"
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CLAUSULA)
        add_run(p, f"CLÁUSULA {romano(num_clausula)} – {titulo_ing}", bold=True)
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)
        artigo = "o " if gen == "m" else "a "
        add_run(p, f"Ingressa na sociedade {artigo}")
        add_run(p, s["nome"].upper(), bold=True)
        add_run(p, ", " + socio_qualificado(s) + ".")
        num_clausula += 1
        # Cláusula de sub-rogação / conhecimento do ingressante
        if gerar_subroga:
            sub_gen = "a" if gen == "f" else "o"
            p = doc.add_paragraph()
            set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CLAUSULA)
            add_run(p, f"CLÁUSULA {romano(num_clausula)} - DECLARAÇÃO DE CONHECIMENTO", bold=True)
            p = doc.add_paragraph()
            set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)
            add_run(p, s["nome"].upper(), bold=True)
            add_run(p, (
                " declara conhecer perfeitamente a situação econômica e financeira da sociedade, "
                "assumindo o ativo e passivo da mesma, ficando dessa forma sub-rogad"
                f"{sub_gen} a todos os direitos e obrigações decorrentes do presente instrumento."
            ))
            num_clausula += 1

    # INGRESSO DE SÓCIOS (direto, sem ser via transferência)
    for novo_s in alteracoes.get("ingresso_socios", []):
        _clausula_ingresso(novo_s)

    # RETIRADA DE SÓCIOS — com destino específico (venda e transferência)
    for ret_s in alteracoes.get("retirada_socios", []):
        gen_r = ret_s.get("genero", "m")
        eh_f  = isinstance(gen_r, str) and gen_r.lower().startswith("f")
        artigo_ret = "a sócia" if eh_f else "o sócio"
        pron_ret   = "A sócia retirante" if eh_f else "O sócio retirante"

        destino_tipo  = ret_s.get("destino_tipo", "socios_remanescentes")
        dest_existente = ret_s.get("destino_socio_existente", "")
        dest_novo_dados = ret_s.get("destino_socio_novo")

        # Se o destino é novo sócio, gera ingresso primeiro
        if destino_tipo == "novo_socio" and dest_novo_dados:
            _clausula_ingresso(dest_novo_dados)

        # Lookup de cotas do retirante nos sócios atuais
        ret_atual = _socio_atual(ret_s.get("nome", ""))
        qtd_ret   = int(ret_atual.get("quantidadeCotas", 0))
        val_ret   = float(ret_atual.get("valorUnitarioCota", 1))
        total_ret = qtd_ret * val_ret

        titulo_ret = "RETIRADA DE SÓCIA" if eh_f else "RETIRADA DE SÓCIO"
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CLAUSULA)
        add_run(p, f"CLÁUSULA {romano(num_clausula)} – {titulo_ret}", bold=True)
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)

        if destino_tipo in ("socio_existente", "novo_socio") and (dest_existente or dest_novo_dados):
            # Formato detalhado: vende e transfere com valor e destinatário
            destinatario = (dest_novo_dados.get("nome","") if dest_novo_dados else dest_existente).upper()
            qtd_ext  = cotas_por_extenso(qtd_ret)
            val_u_ext = valor_por_extenso(val_ret)
            val_tot_ext = valor_por_extenso(total_ret)
            val_tot_fmt = formatar_valor_reais(total_ret)
            artigo_dest = "o sócio" if not (dest_novo_dados and inferir_genero(dest_novo_dados) == "f") else "a sócia"
            qualif_ret = "qualificada" if eh_f else "qualificado"
            possuid_ret = "possuidora" if eh_f else "possuidor"
            add_run(p, f"Retira-se {artigo_ret} ")
            add_run(p, ret_s.get("nome","").upper(), bold=True)
            add_run(p, f", já {qualif_ret}, {possuid_ret} de ")
            add_run(p, f"{qtd_ret:,} ({qtd_ext}) quotas".replace(",","."), bold=True)
            add_run(p, f" no valor de R$ {formatar_valor_reais(val_ret)} ({val_u_ext}) cada, "
                       f"totalmente integralizadas, que vende e transfere, com o consentimento dos sócios, "
                       f"todas suas ")
            add_run(p, f"{qtd_ret:,} ({qtd_ext}) quotas".replace(",","."), bold=True)
            add_run(p, f", pelo valor de R$ {val_tot_fmt} ({val_tot_ext}), para {artigo_dest} ")
            add_run(p, destinatario, bold=True)
            add_run(p, f", acima {'qualificada' if artigo_dest.endswith('a') else 'qualificado'}, "
                       f"dando plena quitação das quotas vendidas, nada mais a reclamar em "
                       f"tempo algum sobre o referido instrumento.")
        else:
            # Formato simples: sócios remanescentes — mas ainda gera quitação plena
            ret_atual2 = _socio_atual(ret_s.get("nome", ""))
            qtd_ret2   = int(ret_atual2.get("quantidadeCotas", 0))
            val_ret2   = float(ret_atual2.get("valorUnitarioCota", 1))
            total_ret2 = qtd_ret2 * val_ret2
            qtd_ext2   = cotas_por_extenso(qtd_ret2) if qtd_ret2 else "todas"
            val_tot_fmt2 = formatar_valor_reais(total_ret2) if total_ret2 else "valor apurado"
            val_tot_ext2 = valor_por_extenso(total_ret2) if total_ret2 else "valor apurado"
            qualif_rem = "qualificada" if eh_f else "qualificado"
            possuid_rem = "possuidora" if eh_f else "possuidor"
            add_run(p, f"Retira-se {artigo_ret} ")
            add_run(p, ret_s.get("nome","").upper(), bold=True)
            add_run(p, f", já {qualif_rem}, {possuid_rem} de ")
            if qtd_ret2:
                add_run(p, f"{qtd_ret2:,} ({qtd_ext2}) quotas".replace(",","."), bold=True)
                add_run(p, f" no valor total de R$ {val_tot_fmt2} ({val_tot_ext2}), totalmente integralizadas, que cede e transfere")
            else:
                add_run(p, "suas quotas")
            add_run(p, ", com o consentimento dos sócios, todas as suas quotas aos demais sócios "
                       "remanescentes, dando plena, geral e irrevogável quitação da sociedade, "
                       "nada mais tendo a reclamar em tempo algum, a qualquer título.")
        num_clausula += 1

    # TRANSFERÊNCIA DE COTAS
    for transf in alteracoes.get("transferencia_cotas", []):
        cedente_nome = transf.get("cedente_nome", "")
        cessionario  = transf.get("cessionario_nome", "").upper()
        ces_tipo     = transf.get("cessionario_tipo", "existente")
        ces_dados    = transf.get("cessionario_dados")
        qtd_t        = int(transf.get("cotas", 0))
        val_t        = float(transf.get("valor", qtd_t))
        qtd_ext      = cotas_por_extenso(qtd_t)
        val_fmt      = formatar_valor_reais(val_t)
        val_ext      = valor_por_extenso(val_t)

        # Lookup do cedente para gênero
        ced_atual = _socio_atual(cedente_nome)
        ced_gen   = inferir_genero(ced_atual) if ced_atual else "m"
        art_ced   = "A sócia remanescente" if ced_gen == "f" else "O sócio remanescente"

        # Se cessionário é novo, gera cláusula de ingresso (sem declaração — virá após a transferência)
        if ces_tipo == "novo" and ces_dados:
            _clausula_ingresso(ces_dados, gerar_subroga=False)
            cessionario = ces_dados.get("nome","").upper()

        ces_gen   = inferir_genero(ces_dados) if ces_dados else "m"
        art_ces   = "à sócia" if ces_gen == "f" else "ao sócio"
        titulo_ces = "ingressante" if ces_tipo == "novo" else "remanescente"

        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CLAUSULA)
        add_run(p, f"CLÁUSULA {romano(num_clausula)} – DA TRANSFERÊNCIA DE QUOTAS", bold=True)
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)
        add_run(p, f"{art_ced} ")
        add_run(p, cedente_nome.upper(), bold=True)
        add_run(p, f" acima qualificad{'a' if ced_gen == 'f' else 'o'}, com o consentimento dos sócios, "
                   f"vende e transfere, ")
        add_run(p, f"{qtd_t:,} ({qtd_ext}) quotas".replace(",","."), bold=True)
        add_run(p, f" pelo valor de R$ {val_fmt} ({val_ext}), para {art_ces} {titulo_ces} ")
        add_run(p, cessionario, bold=True)
        add_run(p, ", dando plena quitação das quotas vendidas.")
        num_clausula += 1

        # Cláusula de declaração de conhecimento — apenas para cessionário novo (ingressante)
        if ces_tipo == "novo":
            p = doc.add_paragraph()
            set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CLAUSULA)
            add_run(p, f"CLÁUSULA {romano(num_clausula)} – DECLARAÇÃO DE CONHECIMENTO", bold=True)
            p = doc.add_paragraph()
            set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)
            art_ces2 = "A sócia" if ces_gen == "f" else "O sócio"
            sub_gen2 = "a" if ces_gen == "f" else "o"
            add_run(p, f"{art_ces2} ingressante ")
            add_run(p, cessionario, bold=True)
            add_run(p, f" declara conhecer perfeitamente a situação econômica e "
                       f"financeira da sociedade, assumindo o ativo e passivo da mesma, ficando "
                       f"dessa forma sub-rogad{sub_gen2} a todos os direitos "
                       f"e obrigações decorrentes do presente instrumento.")
            num_clausula += 1

    # CAPITAL SOCIAL — sempre que houve ingresso, retirada, transferência ou alt. explícita de capital
    houve_mudanca_socios = (
        bool(alteracoes.get("ingresso_socios")) or
        bool(alteracoes.get("retirada_socios")) or
        bool(alteracoes.get("transferencia_cotas")) or
        alteracoes.get("capital_social", {}).get("ativo")
    )
    total_cotas_novas = sum(int(s.get("quantidadeCotas", 0)) for s in socios_novos)
    capital_novo   = empresa_nova.get("capitalSocial", empresa_atual.get("capitalSocial", 0))
    capital_antigo = float(empresa_atual.get("capitalSocial", 0))

    if houve_mudanca_socios and total_cotas_novas > 0:
        cot_fmt = f"{total_cotas_novas:,}".replace(",", ".")
        cot_ext = cotas_por_extenso(total_cotas_novas)
        cap_fmt = formatar_valor_reais(capital_novo)
        cap_ext = valor_por_extenso(capital_novo)

        # CLÁUSULA DE AUMENTO — só quando o capital cresce
        houve_aumento = capital_novo > capital_antigo + 0.01
        if houve_aumento:
            p = doc.add_paragraph()
            set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CLAUSULA)
            add_run(p, f"CLÁUSULA {romano(num_clausula)} – DO AUMENTO DE CAPITAL SOCIAL", bold=True)

            cap_ant_fmt = formatar_valor_reais(capital_antigo)
            cap_ant_ext = valor_por_extenso(capital_antigo)
            p = doc.add_paragraph()
            set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)
            add_run(p, (
                f"O capital social da empresa no valor de R$ {cap_ant_fmt} ({cap_ant_ext}) "
                f"fica elevado para R$ {cap_fmt} ({cap_ext}), divididos em "
                f"{cot_fmt} ({cot_ext}) quotas, no valor de R$ 1,00 (um real) cada uma, "
                f"na forma de seus atos societários e integralizadas da seguinte forma:"
            ))

            # Sócios existentes que aumentaram cotas (seção capital_social ativa)
            socios_atuais_map = {
                s.get("nome", "").upper(): s
                for s in empresa_atual.get("socios", [])
            }
            if alteracoes.get("capital_social", {}).get("ativo"):
                for cs in alteracoes.get("capital_social", {}).get("socios", []):
                    old_s      = socios_atuais_map.get(cs.get("nome", "").upper(), {})
                    old_cotas  = int(old_s.get("quantidadeCotas", 0))
                    new_cotas  = int(cs.get("novas_cotas", old_cotas))
                    val_unit   = float(cs.get("valor_unitario", 1))
                    diff       = new_cotas - old_cotas
                    if diff <= 0:
                        continue
                    valor_sub      = diff * val_unit
                    valor_sub_fmt  = formatar_valor_reais(valor_sub)
                    valor_sub_ext  = valor_por_extenso(valor_sub)
                    gen  = inferir_genero(old_s) if old_s else "m"
                    art  = "A sócia" if gen == "f" else "O sócio"
                    qual = "já qualificada" if gen == "f" else "já qualificado"
                    p = doc.add_paragraph()
                    set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)
                    integ_txt = _texto_integralizacao_partes(cs, valor_sub)
                    add_run(p, f"{art} ")
                    add_run(p, cs.get("nome", "").upper(), bold=True)
                    add_run(p, f", {qual}, subscreve e integraliza neste ato o valor de ")
                    add_run(p, f"R$ {valor_sub_fmt} ({valor_sub_ext})", bold=True)
                    add_run(p, f", {integ_txt}.")

            # Ingressantes que trouxeram capital novo
            for ing in alteracoes.get("ingresso_socios", []):
                qtd_ing    = int(ing.get("quantidadeCotas", 0))
                val_ing    = float(ing.get("valorUnitarioCota", 1))
                valor_sub  = qtd_ing * val_ing
                if valor_sub <= 0:
                    continue
                valor_sub_fmt = formatar_valor_reais(valor_sub)
                valor_sub_ext = valor_por_extenso(valor_sub)
                gen_ing = inferir_genero(ing)
                art_ing = "A sócia ingressante" if gen_ing == "f" else "O sócio ingressante"
                p = doc.add_paragraph()
                set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)
                integ_txt_ing = _texto_integralizacao_partes(ing, valor_sub)
                add_run(p, f"{art_ing} ")
                add_run(p, ing.get("nome", "").upper(), bold=True)
                add_run(p, ", subscreve e integraliza o valor de ")
                add_run(p, f"R$ {valor_sub_fmt} ({valor_sub_ext})", bold=True)
                add_run(p, f", {integ_txt_ing}.")

            num_clausula += 1

        # CLÁUSULA DO CAPITAL SOCIAL — nova distribuição (sempre)
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CLAUSULA)
        add_run(p, f"CLÁUSULA {romano(num_clausula)} - DO CAPITAL SOCIAL", bold=True)
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CAPITAL)
        add_run(p, (
            f"Em razão da presente alteração, o capital social passa a ser de "
            f"R$ {cap_fmt} ({cap_ext}), divididos em {cot_fmt} ({cot_ext}) quotas, "
            f"no valor de R$ 1,00 (um real) cada uma, totalmente subscrito e integralizado."
        ))
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)
        add_run(p, "Parágrafo único.", bold=True)
        cap_pu = _t("alt_capital_pu",
            "O capital encontra-se subscrito e integralizado em moeda corrente nacional "
            "pelos sócios, distribuídos da seguinte forma:")
        add_run(p, f" {cap_pu}")
        gerar_tabela_cotas(doc, socios_novos, total_cotas_novas)
        num_clausula += 1

    # ADMINISTRAÇÃO
    if alteracoes.get("administracao", {}).get("ativo"):
        adm_alt     = alteracoes["administracao"]
        admins_novos = [s for s in socios_novos if s.get("administrador")]
        tipo_adm    = adm_alt.get("tipoAdministracao", "isolada").lower()

        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CLAUSULA)
        add_run(p, f"CLÁUSULA {romano(num_clausula)} - DA ADMINISTRAÇÃO "
                   f"(ART. 997, VI; 1.013, 1.015; 1.064, CC)", bold=True)

        if len(admins_novos) == 1:
            adm    = admins_novos[0]
            gen    = inferir_genero(adm)
            artigo = "pela sócia" if gen == "f" else "pelo sócio"
            p = doc.add_paragraph()
            set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)
            add_run(p, f"A administração da sociedade será exercida {artigo} ")
            add_run(p, adm["nome"].upper(), bold=True)
            add_run(p, " que representará legalmente a sociedade ")
            add_run(p, "ISOLADAMENTE", bold=True)
            add_run(p, " e poderá praticar todo e qualquer ato de gestão pertinente ao objeto social.")
        elif len(admins_novos) > 1:
            if tipo_adm == "conjunta":
                modo_alt      = "CONJUNTAMENTE"
                sep_nomes_alt = " e "
            elif tipo_adm == "isolada_conjunta":
                modo_alt      = "ISOLADAMENTE e/ou CONJUNTAMENTE"
                sep_nomes_alt = " e/ou "
            else:
                modo_alt      = "ISOLADAMENTE"
                sep_nomes_alt = " e/ou "
            p = doc.add_paragraph()
            set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)
            add_run(p, "A administração da sociedade será exercida pelos sócios ")
            for i, adm in enumerate(admins_novos):
                add_run(p, adm["nome"].upper(), bold=True)
                if i < len(admins_novos) - 2:
                    add_run(p, ", ")
                elif i == len(admins_novos) - 2:
                    add_run(p, sep_nomes_alt)
            add_run(p, ", que representarão legalmente a sociedade ")
            add_run(p, modo_alt, bold=True)
            add_run(p, " e poderão praticar todo e qualquer ato de gestão pertinente ao objeto social.")
        num_clausula += 1

    # DECLARAÇÃO DE DESIMPEDIMENTO — sempre após a cláusula de administração
    novos_adms = [s for s in alteracoes.get("ingresso_socios", []) if s.get("administrador")]
    if novos_adms or alteracoes.get("administracao", {}).get("ativo"):
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL,
                             space_before=SPC_NORMAL, space_after=SPC_CLAUSULA)
        add_run(p, f"CLÁUSULA {romano(num_clausula)} - DECLARAÇÃO DE DESIMPEDIMENTO "
                   f"(art. 1.011, § 1° CC e art. 37, II da Lei n° 8.934 de 1994)", bold=True)
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)
        add_run(p, _t("alt_desimpedimento",
            "O(s) Administrador(es) declara(m), sob as penas da lei, de que não está(ão) "
            "impedido(s) de exercer a administração da empresa, por lei especial, ou em virtude "
            "de condenação criminal, ou se encontrar sob os efeitos dela, a pena que vede, ainda "
            "que temporariamente, o acesso a cargos públicos; ou por crime falimentar, de "
            "prevaricação, peita ou suborno, concussão, peculato, ou contra a economia popular, "
            "contra o sistema financeiro nacional, contra as normas de defesa da concorrência, "
            "contra as relações de consumo, a fé pública ou a propriedade."))
        num_clausula += 1

    # OBJETO SOCIAL / ATIVIDADES
    if alteracoes.get("objeto_social", {}).get("ativo"):
        obj_alt = alteracoes["objeto_social"]
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CLAUSULA)
        add_run(p, f"CLÁUSULA {romano(num_clausula)} - DO OBJETO SOCIAL (art. 997, II, CC)", bold=True)
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CLAUSULA)
        add_run(p, "O objeto social passa a ser: ")
        add_run(p, obj_alt.get("novo", "").upper(), bold=True)
        for atv in obj_alt.get("atividades", []):
            desc_raw = atv.get('descricao', '')
            desc_fmt = (desc_raw[0].upper() + desc_raw[1:].lower()) if desc_raw else desc_raw
            p = doc.add_paragraph()
            set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CNAE)
            add_run(p, f"CNAE Nº {atv['cnae']} - {desc_fmt}")
        num_clausula += 1

    # OUTRAS CLÁUSULAS
    for outra in alteracoes.get("outras_clausulas", []):
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CLAUSULA)
        titulo_outra = outra.get("titulo", f"DISPOSIÇÃO ADICIONAL")
        if not titulo_outra.upper().startswith("CLÁUSULA"):
            titulo_outra = f"CLÁUSULA {romano(num_clausula)} - {titulo_outra.upper()}"
        add_run(p, titulo_outra, bold=True)
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)
        add_run(p, outra.get("corpo", ""))
        num_clausula += 1

    # PERMANECEM INALTERADAS
    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CLAUSULA)
    permanece_body = _t("alt_permanece",
        "Permanecem inalteradas as demais cláusulas vigentes que não colidirem com as "
        "disposições do presente instrumento.")
    add_run(p, f"CLÁUSULA {romano(num_clausula)} - {permanece_body}", bold=True)
    num_clausula += 1

    # DA CONSOLIDAÇÃO — encerra o instrumento; o consolidado segue sem quebra de página
    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CLAUSULA)
    consol_body = _t("alt_consolidacao",
        "A vista das modificações ora ajustadas consolida-se o contrato social, que passa a ter "
        "a seguinte redação:")
    add_run(p, f"CLÁUSULA {romano(num_clausula)} - Da Consolidação:", bold=True)
    add_run(p, f" {consol_body}")


# ---------------------------------------------------------------------------
# Gera o contrato consolidado
# ---------------------------------------------------------------------------

def gerar_consolidado(doc, empresa_nova: dict, socios_novos: list,
                      data_fmt: str, data_contrato: datetime,
                      cidade_foro: str, estado_foro: str,
                      socios_retirantes: list = None,
                      nomes_ingressantes: set = None):

    razao_social  = empresa_nova.get("razaoSocial", "").upper()
    objeto_social = empresa_nova.get("objetoSocial", "")
    atividades    = empresa_nova.get("atividades", [])
    end_comercial = empresa_nova.get("enderecoComercial", {})
    capital       = float(empresa_nova.get("capitalSocial", 0))
    total_cotas   = sum(int(s.get("quantidadeCotas", 0)) for s in socios_novos)
    administradores = [s for s in socios_novos if s.get("administrador")]
    classificacao = empresa_nova.get("classificacao", "").lower()
    microempresa  = classificacao in ("me", "epp") or classificacao.startswith("micro")

    tipo_adm = empresa_nova.get("tipoAdministracao", "")
    if not tipo_adm and administradores:
        tipo_adm = administradores[0].get("tipoAdministracao", "isolada")
    tipo_adm = (tipo_adm or "isolada").lower()

    cidade_tc = title_case(cidade_foro)

    # ---- Sócios qualificados ----
    p = doc.add_paragraph()
    set_paragraph_format(p, space_after=SPC_ABERTURA, left_indent=INDENT_NORMAL)
    qtd = len(socios_novos)
    if qtd == 1:
        gen = inferir_genero(socios_novos[0])
        add_run(p, "O sócio abaixo identificado e qualificado:" if gen == "m"
                   else "A sócia abaixo identificada e qualificada:")
    else:
        add_run(p, "Os sócios abaixo identificados e qualificados:")

    for s in socios_novos:
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_SOCIO, space_after=SPC_NORMAL)
        add_run(p, s["nome"].upper(), bold=True)
        add_run(p, ", " + socio_qualificado(s) + ";")

    # Únicos sócios / sede
    sede      = formatar_endereco(end_comercial)
    gens_novos = [inferir_genero(s) for s in socios_novos]
    if qtd == 1:
        comp_cons = "Única sócia componente" if gens_novos[0] == "f" else "Único sócio componente"
    elif all(g == "f" for g in gens_novos):
        comp_cons = "Únicas sócias componentes"
    else:
        comp_cons = "Únicos sócios componentes"
    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=INDENT_SOCIO, space_after=SPC_NORMAL)
    add_run(p, f"{comp_cons} da sociedade empresária limitada, que gira sob o nome de ")
    add_run(p, razao_social, bold=True)
    add_run(p, f", com sede na {sede}.")

    # CLÁUSULA I — NOME
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                         left_indent=INDENT_NORMAL, space_after=SPC_CLAUSULA)
    add_run(p, "CLÁUSULA I - DO NOME EMPRESARIAL (art. 997, II, CC)", bold=True)
    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)
    add_run(p, "A sociedade adota como nome empresarial: ")
    add_run(p, razao_social + ".", bold=True)

    # CLÁUSULA II — SEDE
    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CLAUSULA)
    add_run(p, " CLÁUSULA II - DA SEDE (art. 997, II, CC)", bold=True)
    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=0, first_line_indent=0, space_after=SPC_NORMAL)
    add_run(p, "A sociedade tem sua sede no seguinte endereço: ")
    add_run(p, sede + ";", bold=True)

    # CLÁUSULA III — OBJETO
    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CLAUSULA)
    add_run(p, "CLÁUSULA III - DO OBJETO SOCIAL (art. 997, II, CC)", bold=True)
    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CLAUSULA)
    add_run(p, "A sociedade tem por objeto o exercício das seguintes atividades econômicas: ")
    add_run(p, objeto_social.upper(), bold=True)
    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CLAUSULA)
    add_run(p, "E exercerá as seguintes atividades:")
    for atv in atividades:
        desc_raw = atv.get('descricao', '')
        desc_fmt = (desc_raw[0].upper() + desc_raw[1:].lower()) if desc_raw else desc_raw
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CNAE)
        add_run(p, f"CNAE Nº {atv['cnae']} - {desc_fmt}")

    # CLÁUSULA IV — INÍCIO E PRAZO
    data_inicio = empresa_nova.get("dataInicio", data_contrato.strftime("%d/%m/%Y"))
    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)
    add_run(p, "CLÁUSULA IV - DO INÍCIO DAS ATIVIDADES E PRAZO DE DURAÇÃO "
               "(art. 53, III, F, Decreto n° 1.800/96) ", bold=True)
    add_run(p, "A sociedade iniciou suas atividades em ")
    add_run(p, data_inicio, bold=True)
    add_run(p, " e seu prazo de duração é por tempo indeterminado.")

    # CLÁUSULA V — CAPITAL
    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CLAUSULA)
    add_run(p, "CLÁUSULA V - DO CAPITAL (ART. 997, III e IV e ART. 1.052 e 1.055, CC)", bold=True)

    cap_fmt = formatar_valor_reais(capital)
    cap_ext = valor_por_extenso(capital)
    cot_fmt = f"{total_cotas:,}".replace(",", ".")
    cot_ext = cotas_por_extenso(total_cotas)

    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CAPITAL)
    add_run(p, (
        f"O capital é de R$ {cap_fmt} ({cap_ext}), divididos em {cot_fmt} ({cot_ext}) quotas, "
        f"no valor de R$ 1,00 (um real) cada uma."
    ))
    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)
    add_run(p, "Parágrafo único.", bold=True)
    add_run(p, " O capital encontra-se subscrito e integralizado em moeda corrente nacional, "
               "distribuído pelos sócios da seguinte forma:")
    gerar_tabela_cotas(doc, socios_novos, total_cotas)

    # CLÁUSULA VI — ADMINISTRAÇÃO
    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=INDENT_NORMAL,
                         space_before=SPC_NORMAL, space_after=SPC_CLAUSULA)
    add_run(p, "CLÁUSULA VI - DA ADMINISTRAÇÃO (ART. 997, VI; 1.013, 1.015; 1.064, CC)", bold=True)

    if len(administradores) == 1:
        adm = administradores[0]
        gen = inferir_genero(adm)
        artigo = "pela sócia" if gen == "f" else "pelo sócio"
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CAPITAL)
        add_run(p, f"A administração da sociedade será exercida {artigo} ")
        add_run(p, adm["nome"].upper(), bold=True)
        add_run(p, " que representará legalmente a sociedade ")
        add_run(p, "ISOLADAMENTE", bold=True)
        add_run(p, " e poderá praticar todo e qualquer ato de gestão pertinente ao objeto social.")
    elif len(administradores) > 1:
        if tipo_adm == "conjunta":
            modo_cons      = "CONJUNTAMENTE"
            sep_nomes_cons = " e "
        elif tipo_adm == "isolada_conjunta":
            modo_cons      = "ISOLADAMENTE e/ou CONJUNTAMENTE"
            sep_nomes_cons = " e/ou "
        else:
            modo_cons      = "ISOLADAMENTE"
            sep_nomes_cons = " e/ou "
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CAPITAL)
        add_run(p, "A administração da sociedade será exercida pelos sócios ")
        for i, adm in enumerate(administradores):
            add_run(p, adm["nome"].upper(), bold=True)
            if i < len(administradores) - 2:
                add_run(p, ", ")
            elif i == len(administradores) - 2:
                add_run(p, sep_nomes_cons)
        add_run(p, ", que representarão legalmente a sociedade ")
        add_run(p, modo_cons, bold=True)
        add_run(p, " e poderão praticar todo e qualquer ato de gestão pertinente ao objeto social.")

    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)
    add_run(p, "Parágrafo único.", bold=True)
    add_run(p, " Não constituindo o objeto social, a alienação ou a oneração de bens imóveis "
               "depende de autorização da maioria.")

    # CLÁUSULAS FIXAS VII a IX
    clausulas_fixas = [
        ("CLÁUSULA VII - DO BALANÇO PATRIMONIAL (art. 1.065, CC)",
         "Ao término de cada exercício, em 31 de dezembro, o administrador prestará contas "
         "justificadas de sua administração, procedendo à elaboração do inventário, do balanço "
         "patrimonial e do balanço de resultado econômico."),
        ("CLÁUSULA VIII - DECLARAÇÃO DE DESIMPEDIMENTO DE ADMINISTRADOR "
         "(art. 1.011, § 1° CC e art. 37, II da Lei n° 8.934 de 1994",
         "O Administrador declara, sob as penas da lei, de que não está impedido de exercer "
         "a administração da empresa, por lei especial, ou em virtude de condenação criminal, ou "
         "se encontrar sob os efeitos dela, a pena que vede, ainda que temporariamente, o acesso "
         "a cargos públicos; ou por crime falimentar, de prevaricação, peita ou suborno, "
         "concussão, peculato, ou contra a economia popular, contra o sistema financeiro nacional, "
         "contra as normas de defesa da concorrência, contra as relações de consumo, a fé pública "
         "ou a propriedade."),
        ("CLÁUSULA IX - DO PRÓ LABORE",
         "Os sócios poderão, de comum acordo, fixar uma retirada mensal, a título de pró-labore "
         "para os sócios administradores, observadas as disposições regulamentares pertinentes."),
    ]
    for titulo, corpo in clausulas_fixas:
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CLAUSULA)
        add_run(p, titulo, bold=True)
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)
        add_run(p, corpo)

    # CLÁUSULA X — DISTRIBUIÇÃO DE LUCROS
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                         left_indent=INDENT_NORMAL, space_after=SPC_CLAUSULA)
    add_run(p, "CLÁUSULA X - DISTRIBUIÇÃO DE LUCROS", bold=True)
    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)
    add_run(p, "A sociedade poderá levantar balanços intermediários ou intercalares e "
               "distribuir os lucros evidenciados nos mesmos.")
    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)
    add_run(p, "Parágrafo Primeiro:", bold=True)
    add_run(p, " Os eventuais lucros serão distribuídos entre os sócios, podendo ser "
               "desproporcional aos percentuais de participação societária, conforme deliberação "
               "dos sócios.")
    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)
    add_run(p, "Parágrafo Segundo:", bold=True)
    add_run(p, " Os prejuízos porventura havidos serão transferidos aos exercícios seguintes, "
               "observadas as disposições legais, e suportados pelos sócios na proporção de suas quotas.")

    # CLÁUSULAS XI a XIII + eventual XIV (ME/EPP) + foro
    num_foro = "XV" if microempresa else "XIV"
    clausulas_restantes = [
        ("CLÁUSULA XI - DA RETIRADA OU FALECIMENTO DE SÓCIO",
         "Retirando-se, falecendo ou interditado qualquer sócio, a sociedade continuará suas "
         "atividades com os herdeiros, sucessores e o incapaz, desde que autorizados pelo(s) "
         "outro(s) sócio(s). Não sendo possível ou desejável, a sociedade dissolverá com relação "
         "ao sócio, procedendo-se ao levantamento do balanço especial na data da ocorrência do fato.",
         [("Parágrafo único", " - O mesmo procedimento será adotado em outros casos em que a "
           "sociedade se resolva em relação a seu sócio.")]),
        ("CLÁUSULA XII - DA CESSÃO DE QUOTAS",
         "As quotas são indivisíveis e não poderão ser cedidas ou transferidas a terceiros sem o "
         "consentimento do outro sócio, a quem fica assegurado, em igualdade de condições e preço, "
         "direito de preferência, devendo o cedente comunicar sua intenção ao cessionário, que "
         "terá o prazo de trinta dias para manifestar-se, decorrido o qual, sem manifestação, "
         "entender-se-á que o sócio renunciou ao direito de preferência, ensejando, após a "
         "cessão delas, a alteração contratual pertinente.", []),
        ("CLÁUSULA XIII - DA RESPONSABILIDADE",
         "A responsabilidade de cada sócio é restrita ao valor das suas quotas, mas todos "
         "respondem solidariamente pela integralização do capital social.", []),
    ]

    if microempresa:
        clausulas_restantes.append((
            "CLÁUSULA XIV - PORTE EMPRESARIAL",
            "Os sócios declaram que a sociedade se enquadra como Microempresa - ME, nos termos da "
            "Lei Complementar nº 123, de 14 de dezembro de 2006, e que não se enquadra em qualquer "
            "das hipóteses de exclusão do tratamento favorecido da mesma lei.", []
        ))

    estado_foro_nome = ESTADOS_NOMES.get(estado_foro.upper(), estado_foro.upper())
    clausulas_restantes.append((
        f"CLÁUSULA {num_foro} - DO FORO",
        f"Fica eleito o Foro da Comarca de {cidade_tc} - {estado_foro_nome}, "
        f"com expressa renúncia a qualquer outro, por mais privilegiado que seja.", []
    ))

    for titulo, corpo, paragrafos_unicos in clausulas_restantes:
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_CLAUSULA)
        add_run(p, titulo, bold=True)
        p = doc.add_paragraph()
        set_paragraph_format(p, left_indent=INDENT_NORMAL,
                             space_after=SPC_NORMAL if not paragrafos_unicos else SPC_CAPITAL)
        add_run(p, corpo)
        for bold_part, rest in paragrafos_unicos:
            p = doc.add_paragraph()
            set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)
            add_run(p, bold_part, bold=True)
            add_run(p, rest)

    # Fecho consolidado
    fecho_tpl = _t("alt_fecho",
        "E por estarem em perfeito acordo, em tudo que neste instrumento particular foi lavrado, "
        "obrigam-se a cumprir o presente instrumento, e o assina em uma única via que será "
        "destinada ao registro e arquivamento na Junta Comercial do {ESTADO}.")
    fecho_txt = fecho_tpl.replace("{ESTADO}", estado_foro_nome)
    p = doc.add_paragraph()
    set_paragraph_format(p, left_indent=INDENT_NORMAL, space_after=SPC_NORMAL)
    add_run(p, fecho_txt)

    # Data — cidade/estado da sede da empresa
    cidade_sede = title_case(end_comercial.get("cidade", cidade_foro))
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         left_indent=8890, first_line_indent=0, space_after=SPC_DATA)
    add_run(p, f"{cidade_sede} - {estado_foro_nome}, {data_fmt}.")

    # Assinaturas
    LARGURA_TABELA = TABLE_WIDTH
    largura_col    = LARGURA_TABELA // 2

    # Monta lista de assinantes: remanescentes + ingressantes + retirantes
    _ingressantes_nomes = nomes_ingressantes or set()
    _retirantes         = socios_retirantes or []

    assinantes = []
    for s in socios_novos:
        tipo = "ingressante" if s.get("nome", "").upper() in _ingressantes_nomes else "normal"
        assinantes.append((s, tipo))
    for s in _retirantes:
        assinantes.append((s, "ex"))

    i = 0
    while i < len(assinantes):
        par = assinantes[i:i+2]
        ncols = len(par)
        tbl = doc.add_table(rows=1, cols=ncols)
        remover_bordas_tabela(tbl)
        set_table_width(tbl, LARGURA_TABELA)
        for j, (socio, tipo) in enumerate(par):
            cell = tbl.rows[0].cells[j]
            set_col_width(cell, largura_col)
            bloco_assinatura_cell(cell, socio["nome"], papel_socio(socio, tipo))
        i += 2


# ---------------------------------------------------------------------------
# Função principal
# ---------------------------------------------------------------------------

def gerar_alteracao(dados: dict, caminho_saida):
    """
    dados = {
      "tipo": "alteracao",
      "numero_alteracao": 1,
      "data": "2026-04-26",   (opcional — usa hoje se ausente)
      "empresa_atual": { razaoSocial, cnpj, nire, classificacao,
                         enderecoComercial, objetoSocial, atividades,
                         capitalSocial, socios, dataInicio },
      "alteracoes": {
        "endereco":           { ativo: bool, novo: {...} },
        "ingresso_socios":    [ socio... ],
        "retirada_socios":    [ {nome, cpf, genero}... ],
        "transferencia_cotas": [ {cedente_nome, cessionario_nome, cotas, valor, total}... ],
        "administracao":      { ativo: bool, tipoAdministracao },
        "objeto_social":      { ativo: bool, novo: str, atividades: [...] },
        "outras_clausulas":   [ {titulo, corpo}... ]
      }
    }
    """
    num_alt      = dados.get("numero_alteracao", 1)
    empresa_atual = dados["empresa_atual"]
    alteracoes    = dados.get("alteracoes", {})

    # Data
    data_str = dados.get("data", "")
    try:
        data_contrato = datetime.strptime(data_str, "%Y-%m-%d") if data_str else datetime.today()
    except Exception:
        data_contrato = datetime.today()
    data_fmt = formatar_data_por_extenso(data_contrato)

    # Foro = cidade da sede (atual ou nova)
    if alteracoes.get("endereco", {}).get("ativo"):
        end_foro = alteracoes["endereco"]["novo"]
    else:
        end_foro = empresa_atual.get("enderecoComercial", {})
    cidade_foro = end_foro.get("cidade", "Francisco Beltrão")
    estado_foro = end_foro.get("estado", "PR").upper()

    # Calcula estado NOVO da empresa após as alterações
    empresa_nova = dict(empresa_atual)
    if alteracoes.get("nome_empresarial", {}).get("ativo"):
        empresa_nova["razaoSocial"] = alteracoes["nome_empresarial"].get("novo", "").upper()
    if alteracoes.get("endereco", {}).get("ativo"):
        empresa_nova["enderecoComercial"] = alteracoes["endereco"]["novo"]
    if alteracoes.get("objeto_social", {}).get("ativo"):
        obj_alt = alteracoes["objeto_social"]
        empresa_nova["objetoSocial"] = obj_alt.get("novo", empresa_atual.get("objetoSocial", ""))
        empresa_nova["atividades"]   = obj_alt.get("atividades", empresa_atual.get("atividades", []))

    # Calcula sócios novos
    socios_atuais = list(empresa_atual.get("socios", []))
    cpfs_retirada = {s.get("cpf", "") for s in alteracoes.get("retirada_socios", [])}
    # Cópias mutáveis — redistribuição não altera os dados originais
    socios_novos  = [dict(s) for s in socios_atuais if s.get("cpf", "") not in cpfs_retirada]
    socios_novos += [dict(s) for s in alteracoes.get("ingresso_socios", [])]
    # Adiciona novos sócios que vieram como destino de retirada
    # (usa nome como chave — CPF pode ter sido digitado errado pelo usuário)
    for ret_s in alteracoes.get("retirada_socios", []):
        if ret_s.get("destino_tipo") == "novo_socio" and ret_s.get("destino_socio_novo"):
            dest      = ret_s["destino_socio_novo"]
            dest_nome = dest.get("nome", "").upper()
            if dest_nome and not any(s.get("nome","").upper() == dest_nome for s in socios_novos):
                socios_novos.append(dict(dest))
    # Adiciona cessionários novos de transferência de cotas
    for transf in alteracoes.get("transferencia_cotas", []):
        if transf.get("cessionario_tipo") == "novo" and transf.get("cessionario_dados"):
            ces      = transf["cessionario_dados"]
            ces_nome = ces.get("nome", "").upper()
            if ces_nome and not any(s.get("nome","").upper() == ces_nome for s in socios_novos):
                socios_novos.append(dict(ces))

    # ------------------------------------------------------------------
    # Redistribui cotas — capital total não muda com vendas/transferências
    # (muda apenas se alt_capital estiver ativo)
    # ------------------------------------------------------------------
    def _find_novo(cpf=None, nome=None):
        for s in socios_novos:
            if cpf  and s.get("cpf","")  == cpf:  return s
            if nome and s.get("nome","").upper() == nome.upper(): return s
        return None

    # 1) Retiradas: redistribui cotas do retirante ao destinatário
    for ret_s in alteracoes.get("retirada_socios", []):
        ret_cpf = ret_s.get("cpf", "")
        ret_orig = next(
            (s for s in socios_atuais
             if (ret_cpf and s.get("cpf","") == ret_cpf)
             or s.get("nome","").upper() == ret_s.get("nome","").upper()),
            {}
        )
        qtd_ret = int(ret_orig.get("quantidadeCotas", 0))
        if qtd_ret == 0:
            continue

        destino = ret_s.get("destino_tipo", "socios_remanescentes")

        if destino == "socio_existente":
            dest_nome = ret_s.get("destino_socio_existente", "")
            dest = _find_novo(nome=dest_nome)
            if dest:
                dest["quantidadeCotas"] = int(dest.get("quantidadeCotas", 0)) + qtd_ret

        elif destino == "novo_socio":
            dest_dados = ret_s.get("destino_socio_novo") or {}
            # Busca por nome (CPF pode ter sido digitado errado)
            dest = _find_novo(nome=dest_dados.get("nome", ""))
            if dest and int(dest.get("quantidadeCotas", 0)) == 0:
                dest["quantidadeCotas"] = qtd_ret

        elif destino == "socios_remanescentes":
            remanescentes = [s for s in socios_novos if s.get("cpf","") not in cpfs_retirada]
            total_rem = sum(int(s.get("quantidadeCotas", 0)) for s in remanescentes)
            if len(remanescentes) == 1:
                remanescentes[0]["quantidadeCotas"] = int(remanescentes[0].get("quantidadeCotas", 0)) + qtd_ret
            elif len(remanescentes) > 1 and total_rem > 0:
                distribuido = 0
                for i, s in enumerate(remanescentes):
                    if i < len(remanescentes) - 1:
                        share = round(qtd_ret * int(s.get("quantidadeCotas", 0)) / total_rem)
                        s["quantidadeCotas"] = int(s.get("quantidadeCotas", 0)) + share
                        distribuido += share
                    else:
                        # O último absorve o restante (evita arredondamento)
                        s["quantidadeCotas"] = int(s.get("quantidadeCotas", 0)) + (qtd_ret - distribuido)

    # 2) Transferências parciais: atualiza cedente e cessionário
    for transf in alteracoes.get("transferencia_cotas", []):
        qtd_t    = int(transf.get("cotas", 0))
        ces_tipo = transf.get("cessionario_tipo", "existente")
        if qtd_t == 0:
            continue

        # Subtrai do cedente
        cedente = _find_novo(nome=transf.get("cedente_nome", ""))
        if cedente:
            cedente["quantidadeCotas"] = max(0, int(cedente.get("quantidadeCotas", 0)) - qtd_t)

        # Adiciona ao cessionário
        if ces_tipo == "existente":
            cess = _find_novo(nome=transf.get("cessionario_nome", ""))
            if cess:
                cess["quantidadeCotas"] = int(cess.get("quantidadeCotas", 0)) + qtd_t
        elif ces_tipo == "novo":
            ces_dados = transf.get("cessionario_dados") or {}
            # Busca por nome (CPF pode ter sido digitado errado)
            cess = _find_novo(nome=ces_dados.get("nome", ""))
            if cess and int(cess.get("quantidadeCotas", 0)) == 0:
                cess["quantidadeCotas"] = qtd_t

    # 3) Alteração explícita de capital: substitui cotas dos sócios pelas informadas no formulário
    if alteracoes.get("capital_social", {}).get("ativo"):
        for cs in alteracoes["capital_social"].get("socios", []):
            novas = int(cs.get("novas_cotas", 0))
            val_u = float(cs.get("valor_unitario", 1))
            s = _find_novo(cpf=cs.get("cpf"), nome=cs.get("nome"))
            if s:
                s["quantidadeCotas"] = novas
                if val_u > 0:
                    s["valorUnitarioCota"] = val_u

    # Atualiza administrador se informado
    if alteracoes.get("administracao", {}).get("ativo"):
        tipo_adm = alteracoes["administracao"].get("tipoAdministracao", "isolada")
        empresa_nova["tipoAdministracao"] = tipo_adm
        novos_adm_cpfs = {s.get("cpf") for s in alteracoes["administracao"].get("administradores", [])}
        if novos_adm_cpfs:
            for s in socios_novos:
                s["administrador"]     = s.get("cpf") in novos_adm_cpfs
                s["tipoAdministracao"] = tipo_adm if s["administrador"] else s.get("tipoAdministracao","")

    empresa_nova["socios"] = socios_novos

    # Recalcula capital total
    total_cotas = sum(int(s.get("quantidadeCotas", 0)) for s in socios_novos)
    empresa_nova["capitalSocial"] = total_cotas * float(socios_novos[0].get("valorUnitarioCota", 1)) if socios_novos else empresa_atual.get("capitalSocial", 0)

    # ---- Monta o documento ----
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.88)
        section.bottom_margin = Cm(2.50)
        section.left_margin   = Cm(1.99)
        section.right_margin  = Cm(1.75)

    # Remove parágrafo inicial vazio
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    # Cabeçalho do instrumento de alteração
    adicionar_cabecalho_alteracao(
        doc,
        empresa_atual.get("razaoSocial", ""),
        num_alt,
        empresa_atual.get("cnpj", ""),
        empresa_atual.get("nire", ""),
    )

    # Instrumento de alteração
    gerar_instrumento(doc, dados, empresa_atual, empresa_nova,
                      socios_atuais, socios_novos, alteracoes,
                      num_alt, data_fmt, cidade_foro, estado_foro)

    # Espaço simples entre instrumento e consolidado (sem quebra de página)
    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=0, space_after=SPC_CLAUSULA)
    adicionar_cabecalho_consolidado(
        doc,
        empresa_atual.get("razaoSocial", ""),
        empresa_atual.get("cnpj", ""),
        empresa_atual.get("nire", "")
    )

    # Nomes dos ingressantes (para marcar papel nas assinaturas)
    # Usa nomes em vez de CPFs para evitar colisões por erro de digitação
    nomes_ingressantes = {s.get("nome", "").upper() for s in alteracoes.get("ingresso_socios", []) if s.get("nome")}
    # Também marca como ingressante quem veio como destino de retirada
    for ret_s in alteracoes.get("retirada_socios", []):
        dest = ret_s.get("destino_socio_novo")
        if dest and dest.get("nome"):
            nomes_ingressantes.add(dest["nome"].upper())
    # Também marca cessionários novos de transferência
    for transf in alteracoes.get("transferencia_cotas", []):
        if transf.get("cessionario_tipo") == "novo":
            ces_nome = transf.get("cessionario_dados", {}).get("nome", "")
            if ces_nome:
                nomes_ingressantes.add(ces_nome.upper())

    # Retirantes para assinar como ex-sócios
    socios_retirantes = [s for s in socios_atuais if s.get("cpf", "") in cpfs_retirada]

    # Contrato consolidado
    gerar_consolidado(doc, empresa_nova, socios_novos,
                      data_fmt, data_contrato, cidade_foro, estado_foro,
                      socios_retirantes=socios_retirantes,
                      nomes_ingressantes=nomes_ingressantes)

    doc.save(caminho_saida)
    print(f"✓ Alteração contratual gerada: {caminho_saida}")
